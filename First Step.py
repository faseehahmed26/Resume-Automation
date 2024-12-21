from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_hyperlink(paragraph, text, url):
    """Create a hyperlink in a paragraph."""
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_target(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element and a new w:rPr element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Create the u and color elements
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')

    # Join all the xml elements
    rPr.append(color)
    rPr.append(u)
    new_run.append(rPr)
    
    # Create a new text element
    text_el = OxmlElement('w:t')
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink

def create_resume():
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Name
    name_para = doc.add_paragraph()
    name_run = name_para.add_run('Mohammad Faseeh Ahmed')
    name_run.bold = True
    name_run.font.size = Pt(14)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact Info
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.add_run('mm9314@g.rit.edu | +1 585 202 5217 | ')
    
    # Add hyperlinks
    links = [
        ('LinkedIn', '#'),
        ('Github', '#'),
        ('Portfolio', '#'),
        ('Kaggle', '#'),
        ('Tableau', '#')
    ]
    
    for i, (text, url) in enumerate(links):
        create_hyperlink(contact_para, text, url)
        if i < len(links) - 1:
            contact_para.add_run(' | ')

    # Education Section
    education_heading = doc.add_paragraph()
    education_run = education_heading.add_run('EDUCATION')
    education_run.bold = True
    education_run.font.color.rgb = RGBColor(0, 0, 255)

    # RIT Education
    rit_para = doc.add_paragraph()
    rit_para.add_run('Rochester Institute of Technology, Rochester, NY, M.S in Data Science').bold = True
    rit_para.add_run('\t\t\t\t').bold = True
    rit_para.add_run('Aug 2023 - May 2025').bold = True
    
    coursework_para = doc.add_paragraph()
    coursework_para.add_run('Coursework: ').bold = True
    coursework_para.add_run('Neural Networks, Software Engineering for Data Science, Applied Statistics')
    
    gpa_para = doc.add_paragraph()
    gpa_para.add_run('GPA: ').bold = True
    gpa_para.add_run('3.84/4.00')

    # Add additional sections similarly...
    # Skills section
    skills_heading = doc.add_paragraph()
    skills_run = skills_heading.add_run('SKILLS')
    skills_run.bold = True
    skills_run.font.color.rgb = RGBColor(0, 0, 255)

    # Add skills content
    skills = [
        ('Programming Languages', 'Java, Python, C++, R, JavaScript, Go, Julia, Object Oriented Programming(Python, Java)'),
        ('Frameworks', 'PyTorch, Keras, Scikit, Tensorflow, Groovy, Spark, Flask, React, React-Native, NodeJS'),
        ('Databases', 'SQL, MongoDB, SQLite, MySQL, NoSQL, PostgreSQL, DynamoDB, SAS'),
        ('Technologies', 'Docker, Git, AWS, Azure, GCP, Kafka, JSON, Numpy, Pandas, MLflow, Postman, Tableau, Power BI, MS Excel'),
        ('ML Algorithms/Techniques', 'Regression, Classification, Clustering, Recommender Systems, Deep Learning, NLP, A/B Testing, MLOps, Time Series, Optimization, Exploratory Data Analytics, ETL, Forecasting')
    ]

    for category, items in skills:
        skill_para = doc.add_paragraph()
        skill_para.add_run(f'{category}: ').bold = True
        skill_para.add_run(items)

    # Save the document
    doc.save('resume.docx')

if __name__ == "__main__":
    create_resume()