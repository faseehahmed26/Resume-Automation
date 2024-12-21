from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import os

class DocumentGenerator:
    def __init__(self):
        self.doc = Document()
        self.setup_document()

    def setup_document(self):
        """Setup document margins and sections"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

    def add_section(self, title, content, is_header=False):
        """Add a section to the document with proper formatting"""
        if not is_header:
            self.doc.add_paragraph().add_run().add_break()
            heading = self.doc.add_paragraph()
            run = heading.add_run(title.upper())
            run.bold = True
            run.font.size = Pt(12)
        
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(content)
        run.font.size = Pt(10)
        
        if is_header:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.font.size = Pt(11)

    def generate_resume(self, contact_info, content):
        """Generate the complete resume"""
        # Header section
        name = contact_info.get('name', 'FULL NAME')
        self.add_section(name, '', is_header=True)
        
        # Contact information
        contact_line = " | ".join([
            contact_info.get('email', ''),
            contact_info.get('phone', ''),
            contact_info.get('linkedin', '')
        ])
        self.add_section('', contact_line, is_header=True)
        
        # Main sections
        sections = {
            'EDUCATION': content.get('education', ''),
            'SKILLS': content.get('skills', ''),
            'PROFESSIONAL EXPERIENCE': content.get('experience', ''),
            'PROJECTS': content.get('projects', '')
        }
        
        for title, section_content in sections.items():
            self.add_section(title, section_content)
        
        # Save the document
        output_path = os.path.join('app', 'static', 'generated', 'resume.docx')
        self.doc.save(output_path)
        return output_path