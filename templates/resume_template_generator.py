from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import os

class ResumeTemplateGenerator:
    def __init__(self):
        self.doc = Document()
        self.setup_document()

    def setup_document(self):
        """Setup document margins and sections"""
        # Set margin sizes
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

    def add_heading_style(self, paragraph, text, size=12, bold=True, alignment='LEFT', color=RGBColor(0, 0, 0)):
        """Add and format a heading"""
        paragraph.alignment = getattr(WD_ALIGN_PARAGRAPH, alignment)
        run = paragraph.add_run(text)
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(size)
        font.bold = bold
        font.color.rgb = color

    def create_template(self):
        """Create the complete resume template"""
        # Header section
        name_para = self.doc.add_paragraph()
        self.add_heading_style(name_para, "{{full_name}}", size=16, alignment='CENTER')

        # Contact info
        contact_para = self.doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_info = "{{email}} | {{phone}} | {{linkedin}} | {{github}} | {{portfolio}} | {{kaggle}} | {{tableau}}"
        run = contact_para.add_run(contact_info)
        run.font.size = Pt(10)

        # Education Section
        self.doc.add_paragraph().add_run().add_break()
        edu_heading = self.doc.add_paragraph()
        self.add_heading_style(edu_heading, "EDUCATION", bold=True)
        self.doc.add_paragraph("{{education}}")

        # Skills Section
        self.doc.add_paragraph().add_run().add_break()
        skills_heading = self.doc.add_paragraph()
        self.add_heading_style(skills_heading, "SKILLS", bold=True)
        self.doc.add_paragraph("{{skills}}")

        # Experience Section
        self.doc.add_paragraph().add_run().add_break()
        exp_heading = self.doc.add_paragraph()
        self.add_heading_style(exp_heading, "PROFESSIONAL EXPERIENCE AND INTERNSHIPS", bold=True)
        self.doc.add_paragraph("{{experience}}")

        # Projects Section
        self.doc.add_paragraph().add_run().add_break()
        proj_heading = self.doc.add_paragraph()
        self.add_heading_style(proj_heading, "PROJECTS", bold=True)
        self.doc.add_paragraph("{{projects}}")

        # Research Papers Section
        self.doc.add_paragraph().add_run().add_break()
        research_heading = self.doc.add_paragraph()
        self.add_heading_style(research_heading, "RESEARCH PAPERS", bold=True)
        self.doc.add_paragraph("{{research_papers}}")

    def save_template(self, output_path='templates/resume_template.docx'):
        """Save the template to a file"""
        # Ensure the directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        self.doc.save(output_path)

def generate_template():
    """Function to generate the initial template"""
    generator = ResumeTemplateGenerator()
    generator.create_template()
    generator.save_template()

if __name__ == "__main__":
    generate_template()